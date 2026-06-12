import streamlit as st
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
import datetime
import requests

# Konfigurasi Halaman
st.set_page_config(page_title="LAPER-app", layout="wide")

st.title("LAPER DI BPS KOTA SOLOK")
st.write("LAPER DI BPS KOTA SOLOK (Laporan Perjalanan Dinas BPS Kota Solok) merupakan aplikasi otomasi pembuatan Laporan Perjalanan Dinas Pegawai di BPS Kota Solok")

# URL TEMPLAT GITHUB
GITHUB_TEMPLATE_URL = "https://raw.githubusercontent.com/bps1372/gesit_suratdinas/main/templat.docx"

def load_template_from_github(url):
    try:
        response = requests.get(url)
        response.raise_for_status()
        return io.BytesIO(response.content)
    except Exception as e:
        st.error(f"Gagal mengunduh templat: {e}")
        return None

# --- FUNGSI FORMAT TANGGAL INDONESIA ---
def format_tanggal_indo(date_obj, include_hari=False):
    bulan_indo = {
        1: "Januari", 2: "Februari", 3: "Maret", 4: "April",
        5: "Mei", 6: "Juni", 7: "Juli", 8: "Agustus",
        9: "September", 10: "Oktober", 11: "November", 12: "Desember"
    }
    hari_indo = ["Senin", "Selasa", "Rabu", "Kamis", "Jumat", "Sabtu", "Minggu"]
    
    tanggal = date_obj.day
    bulan = bulan_indo[date_obj.month]
    tahun = date_obj.year
    
    if include_hari:
        nama_hari = hari_indo[date_obj.weekday()]
        return f"{nama_hari}, {tanggal} {bulan} {tahun}"
    return f"{tanggal} {bulan} {tahun}"

# --- FUNGSI PENGGANTI TEKS ---
def replace_text_and_keep_style(paragraph, replacements):
    original_text = paragraph.text
    new_text = original_text
    
    for key, val in replacements.items():
        new_text = new_text.replace(key, str(val))
        
    if new_text != original_text:
        is_bold = False
        font_name = None
        font_size = None
        for run in paragraph.runs:
            if run.text.strip():
                is_bold = run.bold
                font_name = run.font.name
                font_size = run.font.size
                break
        
        paragraph.text = new_text
        
        for run in paragraph.runs:
            run.bold = is_bold
            if font_name:
                run.font.name = font_name
            if font_size:
                run.font.size = font_size

# --- Data Referensi ---
list_nama = [
    "Alfianto, S.Kom, M.Kom", "Ade Hartadi, SE", "Nurhafizah, A.Md., S.M.", "Riki Hidayat, S.M.",
    "Lidya, S.Si., M.T, M.Sc", "Yati Oktrina, SST, M.E.", "Zwesti Permatasari, SST, M.Sc", "Harnofel Putra, SE.",
    "Eko Kurniawan, SE", "Mona Dewi Putri Zain, S.Si.", "Desneli Irma, SST, M.Stat.", "Nike Topia, S.M.",
    "Nurul Jannah Amdayani, S.Stat.", "Riana Octomi Yuanas, S.Si.", "Rifki Hidayat, S.Stat", "Yessi Anita Rahim, S.Tr.Stat.",
    "Jafriadi, S.E.", "Zenitha Shafira Arrasya, S.Tr.Stat.", "Rafy Dwi Pareza, A.Md.Stat.", "Reko Pratama, S.P.",
    "Muhammad Alam, A.Md.", "Riki Jondrizal, A.Md.", "Ayu Wahyuni", "Randa Ilhamsyah", "Lainnya"
]
list_jabatan = [
    "Kepala BPS", "Kepala Subbagian Umum", "Statistisi Ahli Madya", "Statistisi Ahli Muda", "Statistisi Ahli Pertama",
    "Statistis Penyelia", "Statistisi Mahir", "Statistisi Terampil", "Pranata Komputer Ahli Madya", "Pranata Komputer Ahli Muda",
    "Pranata Komputer Ahli Pertama", "APK APBN Ahli Pertama", "APK APBN Ahli Muda", "APK APBN Ahli Madya",
    "Fungsional Umum", "Staf BPS", "Staf Subbagian Umum", "Lainnya"
]
list_golongan = ["IV/b", "IV/a", "III/d", "III/c", "III/b", "III/a", "II/c", "IX", "VII", "V", "Lainnya"]

# --- Form Input ---
st.subheader("1. Informasi Perjalanan")
kegiatan = st.text_input("Kegiatan (ditulis huruf KAPITAL)")
nomor_surat = st.text_input("Nomor dan Tanggal Surat Tugas")
tujuan = st.text_input("Tujuan Perjalanan")

col1, col2, col3 = st.columns(3)
with col1:
    pilihan_nama = st.selectbox("Nama", list_nama)
    nama = st.text_input("Ketik Nama Anda") if pilihan_nama == "Lainnya" else pilihan_nama
with col2:
    pilihan_jabatan = st.selectbox("Jabatan", list_jabatan)
    jabatan = st.text_input("Ketik Jabatan Anda") if pilihan_jabatan == "Lainnya" else pilihan_jabatan
with col3:
    pilihan_golongan = st.selectbox("Pangkat/Golongan", list_golongan)
    golongan = st.text_input("Ketik Pangkat/Golongan Anda") if pilihan_golongan == "Lainnya" else pilihan_golongan

waktu = st.date_input("Waktu Perjalanan", [])

dates = []
if len(waktu) == 2:
    dates = [waktu[0] + datetime.timedelta(days=i) for i in range((waktu[1] - waktu[0]).days + 1)]
elif len(waktu) == 1:
    dates = [waktu[0]]

data_harian = []
if dates:
    st.subheader(f"2. Detail Kegiatan Harian ({len(dates)} Hari)")
    for i, dt in enumerate(dates):
        tanggal_str = format_tanggal_indo(dt, include_hari=True)
        with st.expander(f"Detail - {tanggal_str}", expanded=(i==0)):
            # Fitur Input Jumlah Kegiatan Per Hari
            jml_kegiatan = st.number_input(f"Berapa kegiatan pada {tanggal_str}?", min_value=1, max_value=10, value=1, key=f"jml_{i}")
            
            for j in range(jml_kegiatan):
                if jml_kegiatan > 1:
                    st.markdown(f"**🔹 Kegiatan {j+1}**")
                
                c1, c2 = st.columns(2)
                with c1: jam_mulai = st.text_input("Jam Mulai (cth: 08:00)", key=f"jm_{i}_{j}")
                with c2: jam_akhir = st.text_input("Jam Akhir (cth: 16:00)", key=f"ja_{i}_{j}")
                uraian = st.text_area("Uraian", key=f"ur_{i}_{j}")
                foto = st.file_uploader("Upload Dokumentasi", type=['png', 'jpg', 'jpeg'], key=f"ft_{i}_{j}")
                
                # Jika kegiatan ke-2 dst, kosongkan kolom tanggal agar rapi di tabel
                tgl_tabel = tanggal_str if j == 0 else ""
                
                data_harian.append({
                    "tanggal": tgl_tabel,
                    "jam_mulai": jam_mulai,
                    "jam_akhir": jam_akhir,
                    "uraian": uraian,
                    "foto": foto
                })
                
                if j < jml_kegiatan - 1:
                    st.divider()

st.markdown("---")
if st.button("Generate Laporan", type="primary"):
    if not dates:
        st.error("Pilih waktu perjalanan!")
    else:
        with st.spinner("Memproses dokumen dan menyesuaikan tabel..."):
            template_bytes = load_template_from_github(GITHUB_TEMPLATE_URL)
            if template_bytes:
                try:
                    doc = Document(template_bytes)
                    
                    if len(dates) > 1:
                        waktu_str = f"{format_tanggal_indo(dates[0])} - {format_tanggal_indo(dates[-1])}"
                    else:
                        waktu_str = format_tanggal_indo(dates[0])
                    
                    replacements = {
                        "<kegiatan>": kegiatan.upper() if kegiatan else "", 
                        "<nama>": nama, 
                        "<jabatan>": jabatan,
                        "<golongan>": golongan, 
                        "<nomorsurat>": nomor_surat,
                        "<tujuan>": tujuan, 
                        "<waktu>": waktu_str
                    }
                    
                    for p in doc.paragraphs:
                        replace_text_and_keep_style(p, replacements)
                    
                    if len(doc.tables) > 0:
                        tabel_kegiatan = doc.tables[0]
                        
                        tabel_kegiatan.autofit = False
                        tbl = tabel_kegiatan._tbl
                        tblPr = tbl.tblPr
                        tblLayout = OxmlElement('w:tblLayout')
                        tblLayout.set(qn('w:type'), 'fixed')
                        tblPr.append(tblLayout)
                        
                        row_to_delete = None
                        col_widths = []
                        
                        for r in tabel_kegiatan.rows:
                            if "<haritanggal>" in r.cells[0].text:
                                row_to_delete = r
                                for cell in r.cells:
                                    col_widths.append(cell.width)
                                break
                        
                        saved_font_name = "Arial"
                        saved_font_size = Pt(11)
                        
                        if row_to_delete:
                            if row_to_delete.cells[0].paragraphs[0].runs:
                                run_sampel = row_to_delete.cells[0].paragraphs[0].runs[0]
                                if run_sampel.font.name:
                                    saved_font_name = run_sampel.font.name
                                if run_sampel.font.size:
                                    saved_font_size = run_sampel.font.size
                            
                            tabel_kegiatan._tbl.remove(row_to_delete._tr)
                        
                        for hari in data_harian:
                            row_cells = tabel_kegiatan.add_row().cells
                            jumlah_kolom = len(row_cells)
                            
                            for idx, cell in enumerate(row_cells):
                                if idx < len(col_widths) and col_widths[idx] is not None:
                                    cell.width = col_widths[idx]
                            
                            if jumlah_kolom > 0:
                                p0 = row_cells[0].paragraphs[0]
                                run0 = p0.add_run(hari['tanggal'])
                                run0.font.name, run0.font.size = saved_font_name, saved_font_size
                            
                            if jumlah_kolom > 1:
                                p1 = row_cells[1].paragraphs[0]
                                jam_teks = ""
                                if hari['jam_mulai'] or hari['jam_akhir']:
                                    jam_teks = f"{hari['jam_mulai']} - {hari['jam_akhir']}"
                                run1 = p1.add_run(jam_teks)
                                run1.font.name, run1.font.size = saved_font_name, saved_font_size
                            
                            if jumlah_kolom > 2:
                                p2 = row_cells[2].paragraphs[0]
                                run2 = p2.add_run(hari['uraian'])
                                run2.font.name, run2.font.size = saved_font_name, saved_font_size
                            
                            if hari['foto']:
                                if jumlah_kolom > 3:
                                    p3 = row_cells[3].paragraphs[0]
                                    run3 = p3.add_run()
                                    run3.add_picture(hari['foto'], width=Inches(1.5))
                                elif jumlah_kolom == 3:
                                    p2 = row_cells[2].add_paragraph()
                                    run_pic = p2.add_run()
                                    run_pic.add_picture(hari['foto'], width=Inches(1.5))
                                
                        for row in tabel_kegiatan.rows:
                            for cell in row.cells:
                                for p in cell.paragraphs:
                                    replace_text_and_keep_style(p, replacements)
                    
                    bio = io.BytesIO()
                    doc.save(bio)
                    bio.seek(0)
                    st.success("Laporan Berhasil Dibuat!")
                    st.download_button("📥 Download Laporan Perjalanan Dinas", bio, f"Laporan_Perjalandinas_{nama.replace(' ','_')}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                except Exception as e:
                    st.error(f"Error saat memproses dokumen: {e}")
